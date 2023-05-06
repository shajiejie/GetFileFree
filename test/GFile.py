from bs4 import BeautifulSoup
import requests
import docx

url = 'https://www.example.com/.html'
con = requests.get(url)
con.encoding = 'utf-8'
texts = con.text
result = BeautifulSoup(texts, 'lxml')

div1 = result.find('div', attrs={'class': 'con_article con_main'})
div_zj_list = div1.find_all('p')

# 创建docx文档对象
doc = docx.Document()

# 设置正文字体和字号
font = doc.styles['Normal'].font
font.name = '宋体'
font.size = docx.shared.Pt(10)

# 循环写入p标签内容
for div_zj in div_zj_list:
    doc.add_paragraph(div_zj.text)
    # 设置每个段落的字体和字号
    para = doc.paragraphs[-1]
    para.style.font.name = '宋体'
    para.style.font.size = docx.shared.Pt(10)

# 设置标题字体和字号
doc.add_heading(result.title.string, level=1)
heading = doc.paragraphs[-1]
heading.style.font.name = '宋体'
heading.style.font.size = docx.shared.Pt(16)

# 保存文档
file_path = '/Users/example/Desktop/example.docx'
doc.save(file_path)

print("文件已保存至: ", file_path)
